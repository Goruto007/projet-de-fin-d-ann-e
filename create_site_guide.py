from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("guide_fonctionnement_site.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(31, 58, 95)
TEXT = RGBColor(34, 34, 34)
MUTED = RGBColor(85, 85, 85)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line


def set_style_font(style, name="Calibri", size=11, color=TEXT, bold=False):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold


def add_field_run(paragraph, field_code):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)


def add_page_number_footer(section):
    section.different_first_page_header_footer = True
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(footer_p, before=0, after=0, line=1.0)
    run = footer_p.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    add_field_run(footer_p, " PAGE ")
    for run in footer_p.runs:
        set_run_font(run, size=9, color=MUTED)


def add_title_paragraph(doc, text, size=26, color=NAVY, bold=True, after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=after, line=1.0)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    return p


def add_center_paragraph(doc, text, size=11, color=TEXT, italic=False, bold=False, after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=after, line=1.25)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_body_paragraph(doc, text, size=11, bold=False, italic=False, after=6):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=after, line=1.25)
    run = p.add_run(text)
    set_run_font(run, size=size, color=TEXT, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    p.paragraph_format.keep_with_next = True
    return p


def add_labeled_paragraph(doc, label, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=4, line=1.25)
    label_run = p.add_run(f"{label}. ")
    set_run_font(label_run, size=11, color=DARK_BLUE, bold=True)
    body_run = p.add_run(text)
    set_run_font(body_run, size=11, color=TEXT)
    return p


def build_document():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.45)

    doc.core_properties.title = "Guide simple du fonctionnement du site Chris & Compagnie"
    doc.core_properties.subject = "Explication simple du site web"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "HTML, CSS, JavaScript, boutique, instruments, guide"

    normal = doc.styles["Normal"]
    set_style_font(normal, size=11, color=TEXT, bold=False)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, bold in [
        ("Heading 1", 16, BLUE, True),
        ("Heading 2", 13, BLUE, True),
        ("Heading 3", 12, DARK_BLUE, True),
    ]:
        style = doc.styles[style_name]
        set_style_font(style, size=size, color=color, bold=bold)

    # Cover page
    add_center_paragraph(doc, "Guide simple", size=10.5, color=BLUE, bold=True, after=8)
    add_title_paragraph(doc, "Comprendre le fonctionnement du site Chris & Compagnie", size=25, color=NAVY, after=6)
    add_center_paragraph(doc, "HTML, CSS et JavaScript expliqués sans jargon", size=13, color=MUTED, italic=True, after=18)
    add_center_paragraph(
        doc,
        "Ce document montre comment le site est construit, comment il s'affiche et comment il réagit quand un visiteur clique sur un bouton.",
        size=11,
        color=TEXT,
        after=10,
    )
    add_center_paragraph(
        doc,
        "Public visé : une personne non informaticienne qui veut comprendre le site avec des mots simples.",
        size=10.5,
        color=MUTED,
        after=0,
    )
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # 1. One sentence summary
    add_heading(doc, "Le site en une phrase", level=1)
    add_body_paragraph(
        doc,
        "Chris & Compagnie est une boutique en ligne d'instruments et de services musicaux. Le visiteur arrive sur la vitrine, choisit une catégorie, ouvre un produit, l'ajoute au panier, puis passe au paiement.",
    )
    add_body_paragraph(
        doc,
        "On peut voir le site comme un magasin réel : la page d'accueil est la vitrine, les pages catalogue sont les rayons, le panier est le caddie et la page de paiement est la caisse.",
    )
    add_body_paragraph(
        doc,
        "La page instrument-a-corde.html sert de modèle visuel pour les autres catalogues, ce qui donne un aspect cohérent à l'ensemble du site.",
    )

    # 2. Three layers
    add_heading(doc, "Les trois briques qui font tourner le site", level=1)
    add_body_paragraph(
        doc,
        "Trois langages travaillent ensemble. HTML construit la page, CSS la met en forme et JavaScript lui donne des réactions.",
    )

    add_heading(doc, "HTML : la structure", level=2)
    add_body_paragraph(
        doc,
        "HTML est la charpente de la page. Il dit au navigateur quels éléments afficher : titres, textes, images, menus, boutons, cartes produit et formulaires.",
    )
    add_body_paragraph(
        doc,
        "Dans ce projet, le HTML découpe les pages en zones simples : un en-tête en haut, un menu, un contenu principal, puis un pied de page.",
    )
    add_body_paragraph(
        doc,
        "Si le site était une maison, HTML serait les murs et les pièces.",
    )

    add_heading(doc, "CSS : l'apparence", level=2)
    add_body_paragraph(
        doc,
        "CSS sert à rendre le site agréable à lire. Il choisit les couleurs, les tailles, les marges, la forme des cartes, l'apparence des boutons et la façon dont les colonnes se rangent.",
    )
    add_body_paragraph(
        doc,
        "Le fichier theme.css crée une base commune. Les autres feuilles de style ajoutent des réglages selon la page : accueil, vente, location, panier, paiement, maintenance ou compte utilisateur.",
    )
    add_body_paragraph(
        doc,
        "Le CSS rend aussi le site adaptable aux téléphones. Quand l'écran devient petit, les blocs se mettent les uns sous les autres pour rester lisibles.",
    )

    add_heading(doc, "JavaScript : les réactions", level=2)
    add_body_paragraph(
        doc,
        "JavaScript fait bouger le site. Il réagit aux clics, mémorise ce que le visiteur choisit et met à jour les informations sans obliger à tout recharger.",
    )
    add_body_paragraph(
        doc,
        "Dans vente.js, les boutons Ajouter au panier appellent addToCart(...). Le script enregistre l'article dans la mémoire du navigateur et met à jour le petit badge du panier.",
    )
    add_body_paragraph(
        doc,
        "Dans panier.js, le site relit cette mémoire, affiche la liste des articles, calcule le sous-total, la livraison, la taxe et le total, puis permet de retirer un article.",
    )
    add_body_paragraph(
        doc,
        "Dans paiement.js, le formulaire est contrôlé : les champs sont vérifiés, la commande est sauvegardée, puis l'utilisateur est envoyé vers la page de confirmation.",
    )
    add_body_paragraph(
        doc,
        "Important : le panier n'est pas stocké dans une base de données distante. Il est gardé dans le navigateur grâce à localStorage tant que la commande n'est pas validée.",
    )

    # 3. Pages
    add_heading(doc, "Comment les pages sont organisées", level=1)
    add_labeled_paragraph(
        doc,
        "Accueil",
        "index.html est la porte d'entrée du site. Il présente l'entreprise, les services principaux et les boutons qui envoient vers le reste du site.",
    )
    add_labeled_paragraph(
        doc,
        "Boutique principale",
        "vente.html et instrument-de-musique.html servent à entrer dans l'univers des produits. La première page oriente le visiteur, la seconde présente les grandes familles d'instruments.",
    )
    add_labeled_paragraph(
        doc,
        "Catalogues produits",
        "instrument-a-corde.html, instrument-a-vent.html, percussion.html, instrument-electronique.html et peripherique-audio.html affichent des catalogues détaillés avec des cartes produit et des boutons Ajouter au panier.",
    )
    add_labeled_paragraph(
        doc,
        "Services",
        "location.html sert à la location d'instruments et maintenance.html présente l'entretien ou la réparation.",
    )
    add_labeled_paragraph(
        doc,
        "Achat",
        "panier.html récapitule les choix, paiement.html finalise la commande et confirmation.html remercie l'utilisateur après la validation.",
    )
    add_labeled_paragraph(
        doc,
        "Compte client",
        "page_de_connection.html et utilisateur.html gèrent l'accès au compte et les informations de l'utilisateur.",
    )

    # 4. Visitor journey
    add_heading(doc, "Le parcours d'un visiteur", level=1)
    add_body_paragraph(
        doc,
        "Une personne arrive sur la page d'accueil, clique sur Vente, ouvre une famille d'instruments, choisit un produit, l'ajoute au panier, vérifie le résumé, puis passe au paiement.",
    )
    add_body_paragraph(
        doc,
        "Après validation, la commande est confirmée. Le menu du site sert de fil d'Ariane : il permet de revenir en arrière sans se perdre.",
    )
    add_body_paragraph(
        doc,
        "Le tout est pensé pour être simple : l'utilisateur regarde, choisit, ajoute, paie et reçoit une confirmation.",
    )

    # 5. Glossary
    add_heading(doc, "Glossaire simple", level=1)
    add_labeled_paragraph(
        doc,
        "HTML",
        "la structure du site, comme le squelette d'une maison.",
    )
    add_labeled_paragraph(
        doc,
        "CSS",
        "l'habillage visuel : couleurs, tailles, marges et disposition des blocs.",
    )
    add_labeled_paragraph(
        doc,
        "JavaScript",
        "le comportement : boutons, mises à jour automatiques et petites réactions du site.",
    )
    add_labeled_paragraph(
        doc,
        "localStorage",
        "la petite mémoire du navigateur qui garde le panier pendant la navigation.",
    )
    add_labeled_paragraph(
        doc,
        "Responsive",
        "le fait qu'une page s'adapte aux téléphones, tablettes et écrans d'ordinateur.",
    )

    add_heading(doc, "Ce qu'il faut retenir", level=1)
    add_body_paragraph(
        doc,
        "En résumé, HTML construit, CSS habille et JavaScript anime. Ensemble, ces trois éléments transforment une simple page web en boutique interactive.",
    )
    add_body_paragraph(
        doc,
        "Le site Chris & Compagnie fonctionne donc comme un magasin numérique simple à comprendre : on regarde les rayons, on choisit les articles, on les garde dans le panier et on passe à la caisse.",
    )

    add_page_number_footer(section)
    return doc


def main():
    doc = build_document()
    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
